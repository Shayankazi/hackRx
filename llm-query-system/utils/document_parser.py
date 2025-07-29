import os
import tempfile
import uuid
from typing import List, Dict, Any, Optional, Tuple
import aiofiles
import httpx
import pdfplumber
from docx import Document as DocxDocument
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
try:
    import spacy
except ImportError:
    spacy = None
from models.schemas import DocumentType
from config.settings import settings


class DocumentParser:
    def __init__(self):
        # Load spaCy model for text processing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Fallback if model not installed
            self.nlp = None
            print("Warning: spaCy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
    
    async def download_document(self, url: str) -> str:
        """Download document from URL and return temporary file path"""
        # Validate URL has proper protocol
        if not url.startswith(('http://', 'https://')):
            raise ValueError(f"URL must start with http:// or https://. Got: {url}")
            
        async with httpx.AsyncClient() as client:
            response = await client.get(url)
            response.raise_for_status()
            
            # Create temporary file
            suffix = self._get_file_extension(url)
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            temp_file.write(response.content)
            temp_file.close()
            
            return temp_file.name
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename or URL"""
        return os.path.splitext(filename)[1].lower()
    
    async def parse_document(self, file_path: str, document_type: DocumentType) -> Tuple[str, Dict[str, Any]]:
        """Parse document based on type and return text content and metadata"""
        if document_type == DocumentType.PDF:
            return await self._parse_pdf(file_path)
        elif document_type == DocumentType.DOCX:
            return await self._parse_docx(file_path)
        elif document_type == DocumentType.EMAIL:
            return await self._parse_email(file_path)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
    
    async def _parse_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF document using pdfplumber"""
        full_text = ""
        metadata = {"pages": [], "total_pages": 0}
        
        with pdfplumber.open(file_path) as pdf:
            metadata["total_pages"] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                full_text += f"\n--- Page {page_num} ---\n{page_text}\n"
                
                # Extract page-specific metadata
                page_metadata = {
                    "page_number": page_num,
                    "text_length": len(page_text),
                    "has_tables": len(page.extract_tables()) > 0,
                    "char_count": len(page_text)
                }
                metadata["pages"].append(page_metadata)
        
        return full_text.strip(), metadata
    
    async def _parse_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX document using python-docx"""
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        full_text = "\n".join(paragraphs)
        
        # Extract metadata
        metadata = {
            "paragraph_count": len(paragraphs),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "has_tables": len(doc.tables) > 0,
            "table_count": len(doc.tables)
        }
        
        # Extract table content if present
        if doc.tables:
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                table_texts.append("\n".join(table_data))
            
            full_text += "\n\n--- Tables ---\n" + "\n\n".join(table_texts)
        
        return full_text, metadata
    
    async def _parse_email(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse email file (.eml format)"""
        with open(file_path, 'r', encoding='utf-8') as file:
            email_content = file.read()
        
        msg = email.message_from_string(email_content)
        
        # Extract email headers
        metadata = {
            "subject": msg.get("Subject", ""),
            "from": msg.get("From", ""),
            "to": msg.get("To", ""),
            "date": msg.get("Date", ""),
            "message_id": msg.get("Message-ID", "")
        }
        
        # Extract email body
        body_text = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body_text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
        else:
            body_text = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
        
        # Combine subject and body
        full_text = f"Subject: {metadata['subject']}\n\n{body_text}"
        metadata["body_length"] = len(body_text)
        metadata["word_count"] = len(full_text.split())
        
        return full_text, metadata
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[Dict[str, Any]]:
        """Split text into chunks with optional overlap"""
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        # Use spaCy for intelligent chunking if available
        if self.nlp:
            return self._intelligent_chunking(text, chunk_size, overlap)
        else:
            return self._simple_chunking(text, chunk_size, overlap)
    
    def _intelligent_chunking(self, text: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Intelligent chunking using spaCy sentence boundaries"""
        doc = self.nlp(text)
        sentences = [sent.text.strip() for sent in doc.sents]
        
        chunks = []
        current_chunk = ""
        current_size = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            # If adding this sentence would exceed chunk size, create a new chunk
            if current_size + sentence_length > chunk_size and current_chunk:
                chunks.append({
                    "chunk_index": chunk_index,
                    "text": current_chunk.strip(),
                    "word_count": current_size,
                    "char_count": len(current_chunk)
                })
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, overlap)
                current_chunk = overlap_text + " " + sentence
                current_size = len(current_chunk.split())
                chunk_index += 1
            else:
                current_chunk += " " + sentence
                current_size += sentence_length
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append({
                "chunk_index": chunk_index,
                "text": current_chunk.strip(),
                "word_count": current_size,
                "char_count": len(current_chunk)
            })
        
        return chunks
    
    def _simple_chunking(self, text: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Simple word-based chunking"""
        words = text.split()
        chunks = []
        chunk_index = 0
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            
            chunks.append({
                "chunk_index": chunk_index,
                "text": chunk_text,
                "word_count": len(chunk_words),
                "char_count": len(chunk_text)
            })
            chunk_index += 1
        
        return chunks
    
    def _get_overlap_text(self, text: str, overlap_words: int) -> str:
        """Get the last N words for overlap"""
        words = text.split()
        if len(words) <= overlap_words:
            return text
        return " ".join(words[-overlap_words:])
    
    def extract_key_phrases(self, text: str) -> List[str]:
        """Extract key phrases from text using spaCy"""
        if not self.nlp:
            return []
        
        doc = self.nlp(text)
        
        # Extract noun phrases and named entities
        key_phrases = []
        
        # Noun phrases
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) > 1:  # Multi-word phrases
                key_phrases.append(chunk.text.lower().strip())
        
        # Named entities
        for ent in doc.ents:
            key_phrases.append(ent.text.lower().strip())
        
        # Remove duplicates and return
        return list(set(key_phrases))
    
    def cleanup_temp_file(self, file_path: str):
        """Clean up temporary files"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Warning: Could not delete temporary file {file_path}: {e}")


# Global instance
document_parser = DocumentParser()
